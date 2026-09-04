"""Builds a structured Word (.docx) report from a ClusterProject - no Qt
dependency, so it's testable on its own (inspect the returned Document's
paragraphs/tables directly). Replaces the earlier PDF export: a Word
document is something the recipient can actually edit further (add a
letterhead, trim sections, rebrand for a client) rather than a
print-oriented dead end nobody was going to print anyway.

Section order: Servers -> Storage -> Network -> Cluster (config/
thresholds/DR readiness) -> Virtual Machines. Each inventory section
(Servers/Storage/Network) leads with a small aggregate table, then the
full per-device listing below it - the aggregate answers "how much do I
have", the listing answers "what exactly is it".
"""

import sys
from datetime import datetime

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    _docx_import_error = None
except Exception as _exc:  # noqa: BLE001 - deliberately broad, see message below
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_ORIENT = None
    Pt = None
    RGBColor = None
    Inches = None
    _docx_import_error = _exc

from src.calculations.rack import compute_rack_sizing
from src.calculations.sizing import build_failover_report, build_reports
from src.calculations.thresholds import Status, Thresholds
from src.models.cluster_project import PRIMARY, ClusterProject


def _docx_missing_message() -> str:
    detail = f" (underlying error: {_docx_import_error})" if _docx_import_error else ""
    return (
        "Generating a Word report requires python-docx, and it could not be imported"
        f"{detail}.\n\n"
        f"This app is running from:\n{sys.executable}\n\n"
        "If you already ran 'pip install python-docx', make sure you installed "
        "it into THIS interpreter specifically - e.g. run:\n"
        f'"{sys.executable}" -m pip install python-docx\n'
        "A different install elsewhere (system Python, another venv, VS Code's "
        "default interpreter, etc.) won't be seen by this app if it's not the "
        "same environment.\n\n"
        "Also double-check you installed 'python-docx', not the unrelated, "
        "abandoned PyPI package literally named 'docx' - both import as 'docx' "
        "at runtime, but only python-docx provides Document."
    )

# Mirrors src/gui/widgets/status_badge.py's palette - kept as its own
# copy so this module stays Qt-free (same approach html_report.py used).
# Guarded like the imports above - RGBColor(...) would itself raise at
# module-import time if docx failed to import, defeating the point of a
# soft-fail guard.
if RGBColor is not None:
    _STATUS_COLORS = {
        Status.OK: RGBColor(0x2E, 0x7D, 0x32),
        Status.WARNING: RGBColor(0xED, 0x6C, 0x02),
        Status.CRITICAL: RGBColor(0xC6, 0x28, 0x28),
        Status.UNKNOWN: RGBColor(0x75, 0x75, 0x75),
    }
    _GRAY = RGBColor(0x75, 0x75, 0x75)
    _ORANGE = RGBColor(0xED, 0x6C, 0x02)
    _RED = RGBColor(0xC6, 0x28, 0x28)
    _GREEN = RGBColor(0x2E, 0x7D, 0x32)
else:
    _STATUS_COLORS = {}
    _GRAY = _ORANGE = _RED = _GREEN = None


def _add_table(
    document: Document, headers: list[str], rows: list[list[str]], numbered: bool = True,
    column_weights: list[float] | None = None, total_width_inches: float = 9.5,
    font_size_pt: int | None = None,
) -> None:
    if not rows:
        p = document.add_paragraph("None configured.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _GRAY
        return

    display_headers = (["#"] + headers) if numbered else headers
    table = document.add_table(rows=1, cols=len(display_headers))
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    for i, header in enumerate(display_headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    for row_index, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        display_row = ([str(row_index)] + [str(v) for v in row_data]) if numbered else [str(v) for v in row_data]
        for i, value in enumerate(display_row):
            cells[i].text = value

    if font_size_pt is not None and Pt is not None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size_pt)

    if column_weights is not None and Inches is not None:
        display_weights = ([0.3] + column_weights) if numbered else column_weights
        total_weight = sum(display_weights)
        widths = [Inches(total_width_inches * w / total_weight) for w in display_weights]
        table.autofit = False
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]


def _status_text(status: Status) -> tuple[str, RGBColor]:
    return status.value, _STATUS_COLORS.get(status, _GRAY)


def _add_colored_run(paragraph, text: str, color: RGBColor, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.bold = bold


def _yn(value: bool | None) -> str:
    if value is None:
        return "n/a"
    return "Yes" if value else "No"


def _servers_section(document: Document, project: ClusterProject) -> None:
    document.add_heading("Servers", level=1)

    document.add_heading("Summary", level=2)
    agg_rows = []
    for site in project.site_names:
        site_servers = [s for s in project.servers if s.site == site]
        enabled_servers = [s for s in site_servers if s.enabled]
        agg_rows.append([
            site,
            str(len(enabled_servers)),
            str(project.physical_cores(site)),
            str(sum(s.total_cores for s in enabled_servers)),
            f"{project.physical_ram_gb(site):.0f} GB",
        ])
    _add_table(document, ["Site", "Servers (enabled)", "Effective Cores (HT-adj.)", "Physical Cores", "RAM"], agg_rows)

    document.add_heading("All Servers", level=2)
    rows = []
    for s in project.servers:
        rows.append([
            s.name, s.site, s.vendor or "-", s.model or "-",
            f"{s.sockets}x{s.cores_per_socket}", "Yes" if s.hyperthreading_enabled else "No",
            str(s.effective_cores), f"{s.ram_gb} GB",
            "Enabled" if s.enabled else "Disabled",
        ])
    _add_table(document, ["Name", "Site", "Vendor", "Model", "Sockets x Cores", "HT", "Effective Cores", "RAM", "Status"], rows)


def _storage_section(document: Document, project: ClusterProject) -> None:
    document.add_heading("Storage", level=1)

    document.add_heading("Summary", level=2)
    agg_rows = []
    for site in project.site_names:
        site_storages = [s for s in project.storages if s.site == site]
        agg_rows.append([
            site,
            str(len(site_storages)),
            f"{sum(s.raw_capacity_tb for s in site_storages):.1f} TB",
            f"{sum(s.usable_capacity_tb for s in site_storages):.1f} TB",
        ])
    _add_table(document, ["Site", "Storage Systems", "Raw", "Usable"], agg_rows)

    document.add_heading("All Storage Systems", level=2)
    rows = []
    for s in project.storages:
        rows.append([
            s.name, s.site, s.vendor or "-", s.model or "-",
            f"{s.raw_capacity_tb:.1f} TB", f"{s.usable_capacity_tb:.1f} TB",
            f"{s.raid_overhead_percent:.0f}%",
        ])
    _add_table(document, ["Name", "Site", "Vendor", "Model", "Raw", "Usable", "Overhead"], rows)


def _network_section(document: Document, project: ClusterProject) -> None:
    document.add_heading("Network", level=1)

    document.add_heading("Summary", level=2)
    agg_rows = []
    for site in project.site_names:
        site_switches = [s for s in project.switches if s.site == site]
        site_connections = [
            c for c in project.connections
            if any(sw.uid == c.switch_uid for sw in site_switches)
        ]
        agg_rows.append([site, str(len(site_switches)), str(len(site_connections))])
    _add_table(document, ["Site", "Switches", "Connections (via site switches)"], agg_rows)

    document.add_heading("Switches", level=2)
    switch_rows = [
        [sw.name, sw.site, sw.vendor or "-", sw.model or "-", sw.switch_type, str(sw.total_ports)]
        for sw in project.switches
    ]
    _add_table(document, ["Name", "Site", "Vendor", "Model", "Type", "Total Ports"], switch_rows)

    document.add_heading("Connections", level=2)
    server_names = {s.uid: s.name for s in project.servers}
    switch_names = {s.uid: s.name for s in project.switches}
    storage_names = {s.uid: s.name for s in project.storages}
    conn_rows = []
    for c in project.connections:
        endpoint_a = server_names.get(c.server_uid) or storage_names.get(c.storage_uid) or "-"
        endpoint_b = switch_names.get(c.switch_uid) or storage_names.get(c.storage_uid) or "-"
        conn_rows.append([c.connection_kind, endpoint_a, endpoint_b, c.speed, c.media, c.purpose])
    _add_table(document, ["Type", "Endpoint A", "Endpoint B", "Speed", "Media", "Purpose"], conn_rows)


def _cluster_section(document: Document, project: ClusterProject, thresholds: Thresholds) -> None:
    document.add_heading("Cluster", level=1)

    reports = build_reports(project, thresholds)

    for site in project.site_names:
        report = reports[site]
        document.add_heading(f"{site} Site", level=2)

        if report.server_count == 0 and report.vm_count == 0:
            p = document.add_paragraph(f"Nothing configured at {site} yet.")
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = _GRAY
            continue

        rack = compute_rack_sizing(project, site)
        if rack.is_cloud:
            rack_text = "Cloud (not applicable)"
        elif not rack.rack_units:
            rack_text = "n/a"
        elif rack.capacity_u:
            rack_text = f"{rack.rack_units} / {rack.capacity_u} U, {rack.power_watts:.0f} W"
            if rack.over_capacity:
                rack_text = f"\u26a0 {rack_text} (over capacity)"
        else:
            rack_text = f"{rack.rack_units} U, {rack.power_watts:.0f} W"

        metric_rows = [
            ["Servers", str(report.server_count)],
            ["Physical cores (HT-adj.)", str(report.physical_cores)],
            ["Physical RAM", f"{report.physical_ram_gb:.0f} GB"],
            ["Usable storage", f"{report.usable_storage_gb / 1024:.1f} TB"],
            ["VM count", str(report.vm_count)],
            ["vCPU demand (powered on)", str(report.vcpu_demand)],
            ["RAM demand (powered on)", f"{report.ram_demand_gb:.0f} GB"],
            ["Disk demand (all)", f"{report.disk_demand_gb / 1024:.1f} TB"],
            ["Rack Sizing", rack_text],
        ]
        for metric_name, ratio, status, as_percent in (
            ("CPU oversubscription", report.cpu_ratio, report.cpu_status, False),
            ("RAM utilization", report.ram_ratio, report.ram_status, True),
            ("Storage utilization", report.storage_ratio, report.storage_status, True),
        ):
            if ratio is None:
                metric_rows.append([metric_name, "n/a"])
            else:
                value_text = f"{ratio * 100:.0f}%" if as_percent else f"{ratio:.2f} : 1"
                text, _ = _status_text(status)
                metric_rows.append([metric_name, f"{value_text} ({text})"])
        metric_rows.append(["Survives 1 host failure (N+1)", _yn(report.n_plus_one_ok)])
        if report.n_plus_one_ok is False and report.n_plus_one_check is not None:
            check = report.n_plus_one_check
            shortfalls = []
            if not check.ram_ok:
                shortfalls.append(f"+{check.ram_shortfall_gb:.0f} GB RAM")
            if not check.cpu_ok:
                shortfalls.append(f"+{check.cpu_shortfall_effective_cores:.0f} effective CPU cores")
            if shortfalls:
                metric_rows.append(["N+1 shortfall", f"Would need {' and '.join(shortfalls)}"])

        failover = build_failover_report(project, site, thresholds)
        metric_rows.append(["Failover-assigned VMs", str(failover.assigned_vm_count)])
        metric_rows.append(["Failover Ready", _yn(failover.ready)])
        if failover.assigned_vm_count:
            metric_rows.append(["Failover vCPU demand", str(failover.failover_vcpu_demand)])
            metric_rows.append(["Failover RAM demand", f"{failover.failover_ram_demand_gb:.0f} GB"])
            metric_rows.append(["Failover disk demand", f"{failover.failover_disk_demand_gb / 1024:.1f} TB"])

        _add_table(document, ["Metric", "Value"], metric_rows, numbered=False)

    document.add_heading("Assumptions", level=2)
    assumptions = document.add_paragraph()
    assumptions.add_run(
        f"CPU warning/critical: {thresholds.cpu_warning_ratio:.1f}:1 / {thresholds.cpu_critical_ratio:.1f}:1  \u00b7  "
        f"RAM warning/critical: {thresholds.ram_warning_ratio * 100:.0f}% / {thresholds.ram_critical_ratio * 100:.0f}%  \u00b7  "
        f"Storage warning/critical: {thresholds.storage_warning_ratio * 100:.0f}% / {thresholds.storage_critical_ratio * 100:.0f}%"
    ).italic = True


def _vms_section(document: Document, project: ClusterProject) -> None:
    # Landscape - this many columns don't fit portrait width. VMs is the
    # last section in the report, so there's no need to switch back
    # afterward. python-docx applies a NEW section's orientation to
    # everything from here to the end of the document (or the next
    # section break, if there were one).
    new_section = document.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

    document.add_heading("Virtual Machines", level=1)

    assignments_by_vm: dict[str, list[str]] = {}
    for a in project.failover_assignments:
        assignments_by_vm.setdefault(a.vm_uid, []).append(a.target_site)

    vlan_names = {v.uid: v.name for v in project.vlans}
    storage_names = {s.uid: s.name for s in project.storages}
    cluster_names = {c.uid: c.name for c in project.clusters}
    server_names = {s.uid: s.name for s in project.servers}
    pool_names = {
        pool.uid: pool.name
        for storage in project.storages
        for pool in storage.pools
    }

    document.add_heading("All Virtual Machines", level=2)

    # Core columns always show; optional ones only show if at least one
    # VM actually has data in them - a project that's never touched
    # DR Category/VLAN/Storage Pool/Pinned Server/Notes shouldn't get a
    # report with a dozen "-"-filled columns nobody asked for.
    core_headers = ["Name", "Site", "vCPU", "RAM", "Disk", "Power", "Workload Tier"]
    optional_headers = [
        "DR Category", "IP Address", "OS", "VLAN", "Storage Array",
        "Storage Pool", "Cluster", "Pinned Server", "Failover Sites", "Notes",
    ]

    per_vm_optional_values = []
    for vm in project.vms:
        failover_sites = ", ".join(assignments_by_vm.get(vm.uid, [])) or "-"
        per_vm_optional_values.append([
            vm.dr_category or "-", vm.ip_address or "-", vm.os or "-",
            vlan_names.get(vm.vlan_uid, "-"),
            storage_names.get(vm.storage_uid, "-"),
            pool_names.get(vm.storage_pool_uid, "-"),
            cluster_names.get(vm.cluster_uid, "-"),
            server_names.get(vm.pinned_server_uid, "-"),
            failover_sites,
            vm.notes or "-",
        ])

    kept_optional_indices = [
        i for i in range(len(optional_headers))
        if any(values[i] != "-" for values in per_vm_optional_values)
    ]
    headers = core_headers + [optional_headers[i] for i in kept_optional_indices]

    rows = []
    for vm, optional_values in zip(project.vms, per_vm_optional_values):
        core_values = [
            vm.name, vm.site, str(vm.vcpu), f"{vm.ram_gb:.0f} GB", f"{vm.disk_gb:.0f} GB",
            "On" if vm.powered_on else "Off", vm.workload_tier,
        ]
        rows.append(core_values + [optional_values[i] for i in kept_optional_indices])

    _VM_COLUMN_WEIGHTS = {
        "Name": 1.3, "Site": 0.7, "vCPU": 0.5, "RAM": 0.6, "Disk": 0.6,
        "Power": 0.5, "Workload Tier": 1.2, "DR Category": 0.9, "IP Address": 1.0,
        "OS": 1.2, "VLAN": 0.9, "Storage Array": 1.1, "Storage Pool": 1.1,
        "Cluster": 1.1, "Pinned Server": 1.1, "Failover Sites": 1.1, "Notes": 1.5,
    }
    weights = [_VM_COLUMN_WEIGHTS.get(h, 1.0) for h in headers]
    _add_table(document, headers, rows, column_weights=weights, total_width_inches=10.5, font_size_pt=8)

    document.add_heading("Summary", level=2)
    powered_on_vms = [v for v in project.vms if v.powered_on]
    total_vcpu = sum(v.vcpu for v in powered_on_vms)
    total_ram = sum(v.ram_gb for v in powered_on_vms)
    total_disk_all = sum(v.disk_gb for v in project.vms)
    cpu_ratio = project.cpu_oversubscription_ratio(PRIMARY)
    failover_assigned_count = len({a.vm_uid for a in project.failover_assignments})

    summary_rows = [
        ["VMs", str(len(project.vms))],
        ["vCPU Demand (Powered On)", str(total_vcpu)],
        ["RAM Demand (Powered On)", f"{total_ram:.0f} GB"],
        ["CPU Oversub. (Primary)", f"{cpu_ratio:.2f} : 1" if cpu_ratio is not None else "n/a"],
        ["VM Storage (all disks)", f"{total_disk_all:.0f} GB"],
        ["Failover Assigned", str(failover_assigned_count)],
    ]
    _add_table(document, ["Metric", "Value"], summary_rows, numbered=False)


def _eur(amount: float) -> str:
    return f"\u20ac{amount:,.2f}"


def _pct(value: float | None) -> str:
    return f"{value:.1f}%" if value is not None else "n/a"


def _pricing_section(document: Document, project: ClusterProject) -> None:
    from src.calculations.pricing import (
        compute_equipment_pricing,
        compute_maintenance_status,
    )

    document.add_heading("Pricing", level=1)

    equipment = compute_equipment_pricing(project)

    document.add_heading("Equipment Pricing", level=2)
    equipment_rows = [
        [category, _eur(equipment.by_category.get(category, 0.0))]
        for category in ("Servers", "Storage", "Network", "Backup")
    ]
    equipment_rows.append(["Total", _eur(equipment.total)])
    _add_table(document, ["Category", "Total"], equipment_rows)

    if project.maintenance_items:
        document.add_heading("Licenses, Warranties & Maintenance", level=2)
        statuses = compute_maintenance_status(project)
        status_labels = {
            "expired": "Expired",
            "expiring_soon": "Expiring soon",
            "ok": "OK",
            "unknown": "-",
        }
        item_rows = []
        for s in statuses:
            item = s.item
            item_rows.append([
                item.name or "-", item.category, item.applies_to or "-",
                _eur(item.cost), f"{item.duration_months} mo", item.expiry_date or "-",
                status_labels[s.status],
            ])
        _add_table(
            document,
            ["Name", "Category", "Applies To", "Cost", "Duration", "Expiry Date", "Status"],
            item_rows,
        )

        expired = [s for s in statuses if s.status == "expired"]
        expiring = [s for s in statuses if s.status == "expiring_soon"]
        if expired or expiring:
            p = document.add_paragraph()
            if expired:
                names = ", ".join(s.item.name or "(unnamed)" for s in expired)
                _add_colored_run(p, f"Expired: {names}\n", _RED, bold=True)
            if expiring:
                names = ", ".join(s.item.name or "(unnamed)" for s in expiring)
                _add_colored_run(p, f"Expiring within 90 days: {names}", _ORANGE, bold=True)


def build_docx_report(project: ClusterProject, thresholds: Thresholds, app_version: str = "") -> "Document":
    if Document is None:
        raise ImportError(_docx_missing_message())

    document = Document()

    title = document.add_heading(project.name or "Untitled Project", level=0)

    generated = datetime.now().strftime("%Y-%m-%d %H:%M")
    version_line = f"ClusterSizer {app_version}" if app_version else "ClusterSizer"
    subtitle = document.add_paragraph(f"Generated {generated}  \u00b7  {version_line}")
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.color.rgb = _GRAY

    _servers_section(document, project)
    document.add_page_break()
    _storage_section(document, project)
    document.add_page_break()
    _network_section(document, project)
    document.add_page_break()
    _cluster_section(document, project, thresholds)
    document.add_page_break()
    _pricing_section(document, project)
    document.add_page_break()
    _vms_section(document, project)

    return document
